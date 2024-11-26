from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse
import uvicorn
import tempfile
import os
from docx import Document
from anthropic import Anthropic
import json
import shutil
import time
from starlette.background import BackgroundTask
import logging
import subprocess
from dotenv import load_dotenv
import pypdf
import random
import string
from typing import Optional
import zipfile
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables and configure logging
load_dotenv(override=True)
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

CONFIG = {
    "PATHS": {"TEMPLATES": "templates", "OUTPUT": "output"},
    "PDF_SETTINGS": {"MARGIN": "0.75in", "FONT_SIZE": "10pt", "LINE_SPACING": "1.15"},
}

app = FastAPI()

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Templates configuration
templates = Jinja2Templates(directory="templates")


# Initialize Anthropic client with better error handling
def init_anthropic_client():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        logger.error("ANTHROPIC_API_KEY not found in environment variables")
        raise ValueError("ANTHROPIC_API_KEY not found in environment variables")
    try:
        client = Anthropic(api_key=api_key)
        # Test the client with a simple request
        client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=10,
            messages=[{"role": "user", "content": "test"}],
        )
        return client
    except Exception as e:
        logger.error(f"Failed to initialize Anthropic client: {e}")
        raise ValueError(f"Failed to initialize Anthropic client: {e}")


try:
    client = init_anthropic_client()
except Exception as e:
    logger.error(f"Failed to initialize application: {e}")
    raise


def extract_text(file_path: str, file_type: str) -> str:
    """Extract text from PDF or DOCX files."""
    try:
        if file_type == "pdf":
            with open(file_path, "rb") as file:
                pdf = pypdf.PdfReader(file)
                return " ".join(page.extract_text() for page in pdf.pages)
        else:  # docx
            doc = Document(file_path)
            return " ".join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        raise HTTPException(status_code=400, detail=f"Error extracting text: {str(e)}")


def parse_cv_with_llm(
    cv_text: str, job_description: str, scholar_url: Optional[str] = None, personal_website: Optional[str] = None
) -> dict:
    """Use Claude to parse and enhance CV content."""
    try:
        system_prompt = """You are a CV parsing expert. Your task is to parse the CV text and return ONLY a JSON object.

Important Instructions:

- **Do not alter any personal contact details**, including names, emails, phone numbers, addresses, and **URLs**. These should be extracted exactly as they appear in the CV.
- **Do not generate or infer any new contact information**.
- The JSON MUST be valid and follow this exact structure:

{
    "profile": {
        "name": "string",
        "title": "string",
        "summary": "string",
        "contact": {
            "email": "string",
            "phone": "string",
            "location": "string"
        },
        "links": [
            {
                "platform": "string",  // e.g., "GitHub", "LinkedIn", "Google Scholar"
                "url": "string"
            }
        ]
    },
    "work_experience": [
        {
            "title": "string",
            "company": "string",
            "date": "string",
            "achievements": ["string"]
        }
    ],
    "education": [
        {
            "degree": "string",
            "institution": "string",
            "date": "string",
            "details": ["string"]
        }
    ],
    "skills": [
        {
            "category": "string",
            "items": ["string"]
        }
    ],
    "languages": [
        {
            "language": "string",
            "proficiency": "string"
        }
    ]
}"""

        user_content = f"""Parse this CV and enhance it to align closely with the following job description. Focus on highlighting and emphasizing the relevant skills, experiences, and qualifications that match the job requirements. Use keywords and phrases from the job description where appropriate to showcase the alignment.

**Important Instructions:**

- **Do not alter any personal contact details**, including names, emails, phone numbers, addresses, and **URLs**. These should be extracted exactly as they appear in the CV.
- **Do not generate or infer any new contact information**.
- **Ensure that the skills and experiences required in the job description are clearly reflected in the CV content, based on the information provided in the CV. Do not fabricate or include any information that is not present in the CV.**
- **Rephrase or enhance existing descriptions in the CV to use keywords and phrases from the job description, where applicable, to strengthen the alignment.**
- **Return ONLY the JSON object**, no additional text or explanations.

Job Description:
{job_description}

CV Text:
{cv_text}

Personal Website:
{personal_website if personal_website else "Not provided"}

Google Scholar URL:
{scholar_url if scholar_url else "Not provided"}"""

        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4096,
            temperature=0,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": user_content,
                }
            ],
        )

        # Ensure that response content exists
        if not response.content or not response.content[0].text.strip():
            logger.error("Empty response received from Anthropic API.")
            raise HTTPException(
                status_code=500, detail="Empty response received from language model."
            )

        # Get the response text and parse it
        response_text = response.content[0].text.strip()
        try:
            parsed_response = json.loads(response_text)
            return parsed_response
        except json.JSONDecodeError as json_err:
            logger.error(f"JSON decoding failed: {json_err}")
            logger.error(f"Response Text: {response_text}")
            raise HTTPException(
                status_code=500, detail="Invalid JSON response from language model."
            )

    except Exception as e:
        logger.error(f"Error parsing CV with LLM: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def generate_cover_letter(cv_data: dict, job_description: str) -> str:
    """Generate a cover letter using Claude based on CV and job description."""
    try:
        system_prompt = """You are an expert cover letter writer. Your task is to write a compelling, professional cover letter that demonstrates enthusiasm while maintaining appropriate formality.

Key requirements:
- Use the candidate's actual experience and qualifications from their CV
- Extract and use the company name, position title, and team/department information from the job description
- Only use [HIRING_MANAGER] as a placeholder if no hiring manager name is found in the job description
- If company name or position title isn't clear from the job description, use [COMPANY_NAME] or [POSITION_TITLE] as placeholders
- Align the candidate's experience with the job requirements
- Show genuine enthusiasm and motivation
- Maintain a professional yet engaging tone
- Keep the length to one page
- Structure in proper business letter format with date and contact information
- Focus on specific, relevant achievements
- Avoid clichés and generic statements
- Include specific details about the company and role from the job description to show genuine interest

First, analyze the job description to extract:
1. Company name
2. Position title
3. Team/department name
4. Hiring manager's name (if available)
5. Key company values or mission statements
6. Specific projects or initiatives mentioned"""

        user_content = f"""Write a cover letter for this candidate based on their CV and the job description provided. 

Important guidelines:
- Use the candidate's actual name and contact details from the CV
- Extract and use real company details, position, and team information from the job description
- Only use placeholders ([COMPANY_NAME], [POSITION_TITLE], [HIRING_MANAGER]) if the information cannot be found in the job description
- Reference specific achievements and experiences from their CV that align with the job requirements
- Show how their skills and experience make them an ideal fit for this role
- Demonstrate enthusiasm for the role and company while maintaining professionalism
- Make clear connections between their past experiences and the job requirements
- Format the letter properly with date, addresses, and proper spacing
- Include specific details about the company's projects, values, or initiatives mentioned in the job description

CV Data:
{json.dumps(cv_data, indent=2)}

Job Description:
{job_description}"""

        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4096,
            temperature=0.7,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": user_content,
                }
            ],
        )

        if not response.content or not response.content[0].text.strip():
            raise HTTPException(
                status_code=500, detail="Empty response received from language model."
            )

        return response.content[0].text.strip()

    except Exception as e:
        logger.error(f"Error generating cover letter: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def generate_markdown(cv_data: dict) -> str:
    """Convert structured CV data to markdown format."""
    sections = [
        """---
geometry: margin=0.75in
fontsize: 10pt
linestretch: 1.15
---

"""
    ]

    # Profile Section
    profile = cv_data.get("profile", {})
    sections.append(f"# {profile.get('name', 'Your Name')}")
    sections.append(f"**{profile.get('title', 'Your Title')}**\n")

    # Contact Information
    contact = profile.get("contact", {})
    contact_info = []
    if email := contact.get("email"):
        contact_info.append(email)
    if phone := contact.get("phone"):
        contact_info.append(phone)
    if location := contact.get("location"):
        contact_info.append(location)
    if contact_info:
        sections.append(" • ".join(contact_info) + "\n")

    # Links with URL validation
    links = profile.get("links", [])
    if links:
        link_items = []
        for link in links:
            platform = link.get("platform", "")
            url = link.get("url", "")
            if platform and url:
                # Ensure URLs are properly formatted
                if platform == "LinkedIn" and not url.startswith("https://www.linkedin.com/"):
                    url = f"https://www.linkedin.com/in/{url.split('/')[-1]}"
                elif platform == "GitHub" and not url.startswith("https://github.com/"):
                    url = f"https://github.com/{url.split('/')[-1]}"
                link_items.append(f"[{platform}]({url})")
        if link_items:
            sections.append(" • ".join(link_items) + "\n")

    if summary := profile.get("summary"):
        sections.append(f"{summary}\n")

    sections.append("\\rule{\\linewidth}{0.5pt}\n")

    # Work Experience
    work_experience = cv_data.get("work_experience", [])
    if work_experience:
        sections.append("## Work Experience")
        for experience in work_experience:
            title = experience.get("title", "Job Title")
            company = experience.get("company", "Company Name")
            date = experience.get("date", "Date Range")
            achievements = experience.get("achievements", [])
            sections.append(f"### {title} at {company}")
            sections.append(f"*{date}*")
            for achievement in achievements:
                sections.append(f"- {achievement}")
            sections.append("")  # Add an empty line for spacing

    # Education
    education = cv_data.get("education", [])
    if education:
        sections.append("## Education")
        for edu in education:
            degree = edu.get("degree", "Degree")
            institution = edu.get("institution", "Institution Name")
            date = edu.get("date", "Date Range")
            details = edu.get("details", [])
            sections.append(f"### {degree}, {institution}")
            sections.append(f"*{date}*")
            for detail in details:
                sections.append(f"- {detail}")
            sections.append("")  # Add an empty line for spacing

    # Skills
    skills = cv_data.get("skills", [])
    if skills:
        sections.append("## Skills")
        for skill in skills:
            category = skill.get("category", "Category")
            items = skill.get("items", [])
            sections.append(f"### {category}")
            sections.append(", ".join(items))
            sections.append("")  # Add an empty line for spacing

    # Languages
    languages = cv_data.get("languages", [])
    if languages:
        sections.append("## Languages")
        lang_items = []
        for lang in languages:
            language = lang.get("language", "")
            proficiency = lang.get("proficiency", "")
            if language and proficiency:
                lang_items.append(f"**{language}**: {proficiency}")
        sections.append(", ".join(lang_items))
        sections.append("")  # Add an empty line for spacing

    return "\n".join(sections)


def create_pdf(markdown_content: str, output_path: str) -> None:
    """Convert markdown to PDF using pandoc."""
    try:
        with tempfile.NamedTemporaryFile(suffix='.md', mode='w', delete=False) as temp_md:
            temp_md.write(markdown_content)
            temp_md_path = temp_md.name

        cmd = [
            'pandoc',
            temp_md_path,
            '-o', output_path,
            '--pdf-engine=xelatex',
            '--variable', 'colorlinks=true',
            '--variable', 'urlcolor=blue',
            '--variable', 'linkcolor=blue',
            '--standalone'
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            logger.error(f"Pandoc error: {result.stderr}")
            raise Exception(f"Error producing PDF: {result.stderr}")

    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if 'temp_md_path' in locals():
            os.unlink(temp_md_path)


def generate_random_code(length: int = 6) -> str:
    """Generate a random alphanumeric code."""
    try:
        return "".join(random.choices(string.ascii_uppercase + string.digits, k=length))
    except Exception as e:
        logger.error(f"Error generating random code: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def create_cover_letter_docx(content: str, output_path: str) -> None:
    """Create a cover letter in DOCX format with proper formatting."""
    doc = Document()
    
    # Set up the page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add content with proper formatting
    paragraphs = content.split('\n')
    for p in paragraphs:
        if p.strip():  # Only process non-empty lines
            paragraph = doc.add_paragraph()
            if p.startswith('#'):  # Handle headers
                p = p.lstrip('#').strip()
                paragraph.style = 'Heading 1'
            run = paragraph.add_run(p)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # Save the document
    doc.save(output_path)


def create_cv_docx(cv_data: dict, output_path: str) -> None:
    """Create a CV in DOCX format with proper formatting."""
    doc = Document()
    
    # Set up the page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Profile Section
    profile = cv_data.get("profile", {})
    doc.add_heading(profile.get("name", "Your Name"), 0)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(profile.get("title", "Your Title"))
    title_run.bold = True

    # Contact Information
    contact = profile.get("contact", {})
    contact_info = []
    if email := contact.get("email"):
        contact_info.append(email)
    if phone := contact.get("phone"):
        contact_info.append(phone)
    if location := contact.get("location"):
        contact_info.append(location)
    if contact_info:
        doc.add_paragraph(" • ".join(contact_info))

    # Links
    links = profile.get("links", [])
    if links:
        link_items = []
        for link in links:
            platform = link.get("platform", "")
            url = link.get("url", "")
            if platform and url:
                link_items.append(f"{platform}: {url}")
        if link_items:
            doc.add_paragraph(" • ".join(link_items))

    if summary := profile.get("summary"):
        doc.add_paragraph(summary)

    doc.add_paragraph().add_run("_" * 50)  # Horizontal line

    # Work Experience
    if work_experience := cv_data.get("work_experience", []):
        doc.add_heading("Work Experience", 1)
        for experience in work_experience:
            title = experience.get("title", "Job Title")
            company = experience.get("company", "Company Name")
            date = experience.get("date", "Date Range")
            
            exp_heading = doc.add_heading(level=2)
            exp_heading.add_run(f"{title} at {company}")
            
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(date)
            date_run.italic = True

            for achievement in experience.get("achievements", []):
                doc.add_paragraph(achievement, style='List Bullet')

    # Education
    if education := cv_data.get("education", []):
        doc.add_heading("Education", 1)
        for edu in education:
            degree = edu.get("degree", "Degree")
            institution = edu.get("institution", "Institution Name")
            date = edu.get("date", "Date Range")
            
            edu_heading = doc.add_heading(level=2)
            edu_heading.add_run(f"{degree}, {institution}")
            
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(date)
            date_run.italic = True

            for detail in edu.get("details", []):
                doc.add_paragraph(detail, style='List Bullet')

    # Skills
    if skills := cv_data.get("skills", []):
        doc.add_heading("Skills", 1)
        for skill in skills:
            category = skill.get("category", "Category")
            items = skill.get("items", [])
            
            skill_para = doc.add_paragraph()
            skill_para.add_run(f"{category}: ").bold = True
            skill_para.add_run(", ".join(items))

    # Languages
    if languages := cv_data.get("languages", []):
        doc.add_heading("Languages", 1)
        lang_items = []
        for lang in languages:
            language = lang.get("language", "")
            proficiency = lang.get("proficiency", "")
            if language and proficiency:
                lang_items.append(f"{language}: {proficiency}")
        if lang_items:
            doc.add_paragraph(", ".join(lang_items))

    # Save the document
    doc.save(output_path)


@app.post("/upload")
async def upload_files(
    cv_file: UploadFile = File(...),
    job_description: str = Form(...),
    scholar_url: Optional[str] = Form(None),
    personal_website: Optional[str] = Form(None),
    include_cover_letter: Optional[bool] = Form(False),
):
    """Process uploaded CV and generate optimized version."""
    with tempfile.TemporaryDirectory() as temp_dir:
        logger.info(f"Received scholar_url from frontend: {scholar_url}")

        # Save uploaded file
        file_path = os.path.join(temp_dir, cv_file.filename)
        with open(file_path, "wb") as f:
            content = await cv_file.read()
            f.write(content)

        # Extract text
        file_type = "pdf" if cv_file.filename.lower().endswith(".pdf") else "docx"
        cv_text = extract_text(file_path, file_type)

        # Process with LLM
        cv_data = parse_cv_with_llm(cv_text, job_description, scholar_url, personal_website)

        # Handle Google Scholar URL override
        if scholar_url:
            logger.info(f"Attempting to override with scholar_url: {scholar_url}")
            cv_data["profile"]["links"] = [
                link
                for link in cv_data["profile"]["links"]
                if link["platform"] != "Google Scholar"
            ]
            cv_data["profile"]["links"].append(
                {"platform": "Google Scholar", "url": scholar_url.strip()}
            )

        # Generate markdown for CV
        markdown_content = generate_markdown(cv_data)

        # Generate output directory if it doesn't exist
        output_dir = os.path.join(os.path.dirname(__file__), CONFIG["PATHS"]["OUTPUT"])
        os.makedirs(output_dir, exist_ok=True)

        # Generate random code and timestamp for file names
        random_code = generate_random_code()
        timestamp = int(time.time())

        files_to_return = []
        filenames = []

        # Create CV files (both PDF and DOCX)
        cv_filename_pdf = f"cv_{timestamp}_{random_code}.pdf"
        cv_path_pdf = os.path.join(output_dir, cv_filename_pdf)
        create_pdf(markdown_content, cv_path_pdf)
        files_to_return.append(cv_path_pdf)
        filenames.append(cv_filename_pdf)

        cv_filename_docx = f"cv_{timestamp}_{random_code}.docx"
        cv_path_docx = os.path.join(output_dir, cv_filename_docx)
        create_cv_docx(cv_data, cv_path_docx)
        files_to_return.append(cv_path_docx)
        filenames.append(cv_filename_docx)

        # Generate cover letter if requested
        if include_cover_letter:
            cover_letter_content = generate_cover_letter(cv_data, job_description)
            
            # Create DOCX version of cover letter
            cover_letter_filename_docx = f"cover_letter_{timestamp}_{random_code}.docx"
            cover_letter_path_docx = os.path.join(output_dir, cover_letter_filename_docx)
            create_cover_letter_docx(cover_letter_content, cover_letter_path_docx)
            files_to_return.append(cover_letter_path_docx)
            filenames.append(cover_letter_filename_docx)

            # Create PDF version of cover letter
            cover_letter_filename_pdf = f"cover_letter_{timestamp}_{random_code}.pdf"
            cover_letter_path_pdf = os.path.join(output_dir, cover_letter_filename_pdf)
            create_pdf(cover_letter_content, cover_letter_path_pdf)
            files_to_return.append(cover_letter_path_pdf)
            filenames.append(cover_letter_filename_pdf)

        # Create ZIP file with all documents
        zip_filename = f"application_{timestamp}_{random_code}.zip"
        zip_path = os.path.join(output_dir, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file_path, filename in zip(files_to_return, filenames):
                zipf.write(file_path, filename)

        # Clean up individual files
        for file_path in files_to_return:
            os.remove(file_path)

        headers = {
            "Content-Disposition": f'attachment; filename="{zip_filename}"',
            "Content-Type": "application/zip",
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        }

        return FileResponse(
            path=zip_path,
            headers=headers,
            background=BackgroundTask(lambda: os.remove(zip_path)),
        )


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/test")
async def test():
    return {"message": "API is working"}


@app.get("/debug")
async def debug(request: Request):
    try:
        # Try to load and return the raw template content
        template = templates.get_template("index.html")
        content = template.render({"request": request})
        return HTMLResponse(content=content)
    except Exception as e:
        return {"error": str(e), "type": str(type(e))}


if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
